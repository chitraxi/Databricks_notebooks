# Databricks notebook source
# Install the python-docx library
%pip install python-docx

# COMMAND ----------

# Import necessary libraries
from pyspark.sql import functions as F
from pyspark.sql import SparkSession
import pyspark.sql.types as T
from pyspark.sql.window import Window as W
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
import docx.shared
from datetime import datetime
import shutil
from IPython.display import HTML

# Create Spark session
spark = SparkSession.builder.appName("Optimizer Apollo SR insights script").getOrCreate()

# Define date ranges
start_date = '2024-08-01'
end_date = '2024-08-31'
previous_start_date = '2024-07-01'
previous_end_date = '2024-07-31'

# Load data into DataFrame
df_base_data = spark.sql(f"""
    SELECT
        a.merchant_id, 
        COALESCE(a.dba, b.billing_label, b.name) AS name,
        CASE 
            WHEN a.internal_external_flag = 'external' THEN gateway
            WHEN a.internal_external_flag = 'internal' THEN 'Razorpay'
            ELSE NULL
        END AS aggregator,
        a.method_advanced,
        a.network,
        a.network_tokenised_payment,
        IF(a.internal_error_code = '' OR a.internal_error_code IS NULL, 'Success', a.internal_error_code) AS error_code,
        a.wallet,
        a.bank,
        CASE
            WHEN a.internal_error_code IN (
                'list of error codes'
            ) THEN 'customer'
            WHEN a.internal_error_code IN (
                'list of error codes'
            ) THEN 'internal'
            WHEN a.internal_error_code IN (
                'list of error codes'
            ) THEN 'gateway'
            ELSE 'business'
        END AS error_source,
        a.internal_error_code,
        a.created_date,
        CASE 
            WHEN cast(a.created_date as date) BETWEEN date '{start_date}' AND date '{end_date}' THEN 'T-1'
            WHEN cast(a.created_date as date) BETWEEN date '{previous_start_date}' AND date '{previous_end_date}' THEN 'T-2'
            ELSE NULL
        END AS period,
        COUNT(DISTINCT CASE WHEN authorized_at IS NOT NULL or authorized_at != 0 THEN a.id ELSE NULL END) AS payment_success,
        COUNT(DISTINCT a.id) AS payment_attempts
    FROM 
        (
            SELECT * 
            FROM aggregate_ba.payments_optimizer_flag_sincejan2021 
            WHERE created_date BETWEEN '{previous_start_date}' AND '{end_date}' 
            AND optimizer_flag = 1 
            AND golive_date IS NOT NULL 
            AND gateway IS NOT NULL 
            AND method_advanced NOT IN ("Unknown Card", "upi", "card", "paylater", "app", "UPI Unknown","app","intl_bank_transfer")
            and created_date >= golive_date
        ) a 
        INNER JOIN realtime_hudi_api.merchants b 
        ON a.merchant_id = b.id 
        --where a.id not in (select id from realtime_hudi_api.payments where created_date between '{previous_start_date}' AND '{end_date}' and cps_route = 100)
    GROUP BY 
        1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12
""")



#


# COMMAND ----------

# Convert Spark DataFrame to Pandas DataFrame
df_data = df_base_data.fillna('na').toPandas()



# COMMAND ----------



# Create drilled_SR DataFrame
drilled_SR = df_data.pivot_table(
    index=["method_advanced", "aggregator", "merchant_id", "name", "network", "network_tokenised_payment", "bank"],
    columns="period",
    values=["payment_attempts", "payment_success"],
    aggfunc="sum"
).fillna(0).reset_index()

drilled_SR

print(drilled_SR.columns)

# Calculate success rates for T-2 and T-1
drilled_SR["T-2 SR%"] = round((drilled_SR[("payment_success", "T-2")] / drilled_SR[("payment_attempts", "T-2")]) * 100, 2)
drilled_SR["T-1 SR%"] = round((drilled_SR[("payment_success", "T-1")] / drilled_SR[("payment_attempts", "T-1")]) * 100, 2)

# Create SR DataFrame
SR = df_data.pivot_table(
    index=["method_advanced"],
    columns="period",
    values=["payment_success", "payment_attempts"],
    aggfunc=sum
).fillna(0).reset_index()

# Calculate Success Rates for T-2 and T-1
SR["T-2_SR%_agg"] = round((SR[("payment_success", "T-2")] / SR[("payment_attempts", "T-2")])*100 , 4)
SR["T-1_SR%_agg"] = round((SR[("payment_success", "T-1")] / SR[("payment_attempts", "T-1")])*100 , 2)
SR["total_agg"] = SR[("payment_attempts", "T-1")] + SR[("payment_attempts", "T-2")]
SR["Delta_SR_agg"] = round(SR["T-1_SR%_agg"] - SR["T-2_SR%_agg"] , 2)

# Filter for Delta SR < -1 and total > 1000
SR_filtered = SR
#SR[(SR["Delta_SR_agg"] < -1) & (SR["total_agg"] > 3000)]



# COMMAND ----------

SR

# COMMAND ----------

SR_filtered = SR
SR_filtered

# COMMAND ----------

# Merge drilled_SR with SR DataFrame based on index
merged_df = drilled_SR.merge(SR_filtered, on=['method_advanced'], how="inner")
print(merged_df.columns)

# COMMAND ----------

merged_df

# COMMAND ----------

merged_df.to_csv("/dbfs/FileStore/Chitraxi/merged_df.csv")


spark.read.format('csv').load('dbfs:/FileStore/Chitraxi/merged_df.csv').display()

# COMMAND ----------


# Identify rows where payment_attempts_x for T-1 is zero
zero_attempts_mask = merged_df[("payment_attempts_x", "T-1")] == 0
merged_df.loc[zero_attempts_mask, ("T-1 SR%", "")] = merged_df.loc[zero_attempts_mask, ("T-2 SR%", "")]

# Identify rows where payment_attempts_x for T-2 is zero
zero_attempts_mask = merged_df[("payment_attempts_x", "T-2")] == 0
merged_df.loc[zero_attempts_mask, ("T-2 SR%", "")] = merged_df.loc[zero_attempts_mask, ("T-1 SR%", "")]

# Calculate %total_T-1 and %total_T-2
merged_df["%total_T-1"] = round((merged_df[("payment_attempts_x", "T-1")] / merged_df[("payment_attempts_y", "T-1")]), 4)
merged_df["%total_T-2"] = round((merged_df[("payment_attempts_x", "T-2")] / merged_df[("payment_attempts_y", "T-2")]), 4)

# Calculate VMM_SR and VMM_TX
merged_df["VMM_SR"] = round((merged_df["%total_T-2"] * (merged_df[("T-1 SR%", "")] - merged_df[("T-2 SR%", "")])), 2)
merged_df["VMM_TX"] = round((merged_df[("T-1 SR%", "")] * (merged_df["%total_T-1"] - merged_df["%total_T-2"])), 2)




# COMMAND ----------



merged_df.to_csv("/dbfs/FileStore/Chitraxi/merged_df.csv")


spark.read.format('csv').load('dbfs:/FileStore/Chitraxi/merged_df.csv').display()

# COMMAND ----------

merged_df.columns

# COMMAND ----------

# Group by 'method_advanced' and sum 'VMM_SR' and 'VMM_TX'
grouped_df = merged_df.groupby(('method_advanced', '')).agg({('VMM_SR', ''): 'sum', ('VMM_TX', ''): 'sum'}).reset_index()

grouped_df

# COMMAND ----------

method_level_vmm = grouped_df.merge(SR_filtered, on=['method_advanced'], how="inner").reset_index()

method_level_vmm

# COMMAND ----------

# Group by 'method_advanced' and 'aggregator' and select the smallest 5 records based on 'VMM_SR'
top5_vmm_sr = merged_df.groupby(["method_advanced"]).apply(lambda x: x.nlargest(5, "VMM_SR")).reset_index(drop=True)

# COMMAND ----------



top5_vmm_sr.to_csv("/dbfs/FileStore/Chitraxi/top5_vmm_sr.csv")


spark.read.format('csv').load('dbfs:/FileStore/Chitraxi/top5_vmm_sr.csv').display()

# COMMAND ----------

top5_vmm_sr

# COMMAND ----------

top5_vmm_sr = top5_vmm_sr[top5_vmm_sr["VMM_SR"] < -0.8]

# COMMAND ----------

top5_vmm_tx = merged_df.groupby(["method_advanced"]).apply(lambda x: x.nlargest(5, "VMM_TX")).reset_index(drop=True)

#top5_vmm_tx = top5_vmm_tx[top5_vmm_tx["VMM_TX"] < -0.8]


# COMMAND ----------

top5_vmm_tx

# COMMAND ----------



top5_vmm_sr.to_csv("/dbfs/FileStore/Chitraxi/top5_vmm_sr.csv")


spark.read.format('csv').load('dbfs:/FileStore/Chitraxi/top5_vmm_sr.csv').display()

# COMMAND ----------

df_failed = df_data[(df_data["payment_success"] == 0)]

# COMMAND ----------

df_failed

# COMMAND ----------

# Create drilled_SR DataFrame
error_df = df_failed.pivot_table(
    index=["method_advanced", "aggregator", "merchant_id", "name", "network", "network_tokenised_payment", "bank","internal_error_code"],
    columns="period",
    values=["payment_attempts", "payment_success"],
    aggfunc="sum"
).fillna(0).reset_index()

# COMMAND ----------

# Step 1: Confirm column existence
if (("payment_attempts", "T-1") in error_df.columns) and (("payment_attempts", "T-2") in error_df.columns):
    # Step 2: Define the filter condition
    condition = error_df[("payment_attempts", "T-2")] > 0

    # Step 3: Create new column
    error_df["delta_error_count"] = error_df[("payment_attempts", "T-1")] - error_df[("payment_attempts", "T-2")]

    # Step 4: Apply the filter
    error_df = error_df[condition]
else:
    print("Required columns not found in the DataFrame.")

# COMMAND ----------

# Group by the specified columns and sort within each group by payment_attempts
sorted_error_df = error_df.groupby(["method_advanced", "aggregator", "merchant_id", "name", "network", "network_tokenised_payment", "bank"], as_index=False).apply(lambda x: x.sort_values(by=("payment_attempts","T-1"), ascending=False))

# Reset the index
sorted_error_df = sorted_error_df.reset_index(drop=True)

# Show the sorted DataFrame
print(sorted_error_df)


# COMMAND ----------

sorted_error_df

# COMMAND ----------

# Group by specified columns and get top three rows for each group
limited_error_df = sorted_error_df.groupby(["method_advanced", "aggregator", "merchant_id", "name", "network", "network_tokenised_payment", "bank"]).head(3)

# Show the limited DataFrame
limited_error_df

# COMMAND ----------

# Perform inner join on all common columns
top5_vmmSR_error_df = top5_vmm_sr.merge(limited_error_df, on=["method_advanced", "aggregator", "merchant_id", "name", "network", "network_tokenised_payment", "bank"], how="inner")



# COMMAND ----------

top5_vmmSR_error_df

top5_vmmSR_error_df["Delta Errors"] = round(((top5_vmmSR_error_df[("payment_attempts", "T-1")] - top5_vmmSR_error_df[("payment_attempts", "T-2")])/ top5_vmmSR_error_df[("payment_attempts", "T-2")] )*100 , 2)

top5_vmmSR_error_df
summary_error_df = top5_vmmSR_error_df[top5_vmmSR_error_df[("Delta Errors", "")] > 0]

summary_error_df

# COMMAND ----------

overall_SR

# COMMAND ----------

# Calculate success rate (SR) for different periods
overall_SR = df_data.pivot_table(
    columns="period",
    values=["payment_attempts", "payment_success"],
    aggfunc="sum"
).fillna(0)

# Calculate SR for different periods
# Assuming overall_SR is your DataFrame
overall_SR['SR_T-1'] = (overall_SR['T-1']['payment_success'] / overall_SR['T-1']['payment_attempts']) * 100
overall_SR['SR_T-2'] = (overall_SR['T-2']['payment_success'] / overall_SR['T-2']['payment_attempts']) * 100
# Fill NaN values with 0
overall_SR = overall_SR.fillna(0)

overall_SR['delta_SR'] = round(overall_SR['SR_T-1'] - overall_SR['SR_T-2'],2)


# Calculate success rate (SR) for different periods
method_SR = df_data.pivot_table(
    index = "method_advanced",
    columns="period",
    values=["payment_attempts", "payment_success"],
    aggfunc="sum"
).fillna(0).reset_index()


# Assuming overall_SR is your DataFrame
method_SR['SR_T-1'] = round((method_SR['payment_success']['T-1'] / method_SR['payment_attempts']['T-1']) * 100,2)
method_SR['SR_T-2'] = round((method_SR['payment_success']['T-2'] / method_SR['payment_attempts']['T-2']) * 100,2)
# Fill NaN values with 0
method_SR = method_SR.fillna(0)

method_SR['delta_SR'] = round(method_SR['SR_T-1'] - method_SR['SR_T-2'],2)


#print("PFB Method level SR delta:")
#display()

# COMMAND ----------

overall_SR

# COMMAND ----------

method_level_vmm

# COMMAND ----------

# Iterate over each row in the DataFrame
for index, row in overall_SR.iterrows():
    delta_SR_overall = round(row["delta_SR"],2)
    SR_T1_overall = round(row["SR_T-1"],2)
    SR_T2_overall = round(row["SR_T-2"],2)

    # Print the values
    print("\n" + "#" * 30)  # Large heading
    print(f"Overall Delata SR: {delta_SR_overall} Previous Month SR:{SR_T2_overall}   Current Month SR: {SR_T1_overall}")
    print("#" * 30 + "\n")     

    break


print("Below are the methodsXgateway where SR has dropped:")

for index, row in method_level_vmm.iterrows():
    method = row[("method_advanced", "")]
    #aggregator = row[("aggregator", "")]
    Delta_SR = row[("Delta_SR_agg", "")]
    VMM_SR = row[("VMM_SR", "")]
    VMM_TX = row[("VMM_TX", "")]
    #print(f"{method} - {aggregator} : SR dropped by: {Delta_SR}")   
    



for index, row in SR_filtered.iterrows():
    method = row[("method_advanced", "")]
    #aggregator = row[("aggregator", "")]
    Delta_SR = row[("Delta_SR_agg", "")]
    print("\n" + "#" * 30)  # Large heading
    print(f"{method}  : SR dropped by: {Delta_SR}")   
    print("#" * 30 + "\n")     

    # Filter top5_vmm_SR_list based on method and aggregator
    filtered_top5_vmm_SR_list = top5_vmm_sr[
        (top5_vmm_sr[("method_advanced", "")] == method)   
    ].iterrows()
    print("#" * 20)  # Medium heading
    print("SR drop due to change in error trends:")
    print("#" * 20 + "\n")  # Medium heading
    Counter1 = 1

    for inner_index, inner_row in filtered_top5_vmm_SR_list:
        # Extract information for the summary
        
        print()
        inner_method = inner_row[("method_advanced", "")]
        inner_aggregator = inner_row[("aggregator", "")]
        merchant_id = inner_row[("merchant_id", "")]
        name = inner_row[("name", "")]
        bank = inner_row[("bank", "")]
        vmm_sr = inner_row[("VMM_SR", "")]
        network = inner_row[("network", "")]
        network_tokenised_payment = inner_row[("network_tokenised_payment", "")]

        if method in ("Credit Card", "Debit Card", "CC emi", "DC emi"):
            print(f"    {Counter1}. For {inner_method} - {inner_aggregator}, {name} - {merchant_id}, network_tokenised_payment - {network_tokenised_payment}, {network}, VMM_sr dropped by {vmm_sr}")

        if method in ("UPI Intent", "UPI Collect"):
            print(f"    {Counter1}. For {inner_method} - {inner_aggregator}, {name} - {merchant_id}, VMM_sr dropped by {vmm_sr}")

        if method in ("netbanking"):
            print(f"    {Counter1}. For {inner_method} - {inner_aggregator}, {name} - {merchant_id}, {bank}, VMM_sr dropped by {vmm_sr}")

        if method in ("wallet"):
            print(f"    {Counter1}. For {inner_method} - {inner_aggregator}, {name} - {merchant_id}, VMM_sr dropped by {vmm_sr}")
        
        filtered_errors = summary_error_df[
        (summary_error_df[("method_advanced", "")] == method) &
        (summary_error_df[("aggregator", "")] == inner_aggregator) & (summary_error_df[("merchant_id", "")] == merchant_id) 
        & (summary_error_df[("bank", "")] == bank) & (summary_error_df[("network", "")] == network) & (summary_error_df[("network_tokenised_payment", "")] == network_tokenised_payment) 
        ].iterrows()
        
        for index,error_row in filtered_errors:
            error = error_row[("internal_error_code", "")]
            delta_error = round(error_row["Delta Errors",""],2)
            current_attempts = error_row[("payment_attempts", "T-1")]
            previous_attempts = error_row[("payment_attempts", "T-2")]
            print(f"        {error} - incresed by {delta_error}. Previous error count - {previous_attempts} , Current error count- {current_attempts}")
    Counter1 +=1

    print("#" * 20)  # Medium heading
    print("SR drop due to change in Volume:")
    print("#" * 20 + "\n")  # Medium heading
    filtered_top5_vmm_tx_list = top5_vmm_tx[
        (top5_vmm_tx[("method_advanced", "")] == method) 
    ].iterrows()

    
    Counter2 = 1 

    for inner_index, inner_row in filtered_top5_vmm_tx_list:
        # Extract information for the summary
        inner_method = inner_row[("method_advanced", "")]
        inner_aggregator = inner_row[("aggregator", "")]
        merchant_id = inner_row[("merchant_id", "")]
        name = inner_row[("name", "")]
        bank = inner_row[("bank", "")]
        vmm_tx = inner_row[("VMM_TX", "")]
        network = inner_row[("network", "")]
        network_tokenised_payment = inner_row[("network_tokenised_payment", "")]
        current_attempts = inner_row[("payment_attempts_x", "T-1")]
        previous_attempts = inner_row[("payment_attempts_x", "T-2")]
        try:
            delta_traffic = round(((inner_row[("payment_attempts_x", "T-1")] - inner_row[("payment_attempts_x", "T-2")]) / inner_row[("payment_attempts_x", "T-2")]) * 100, 2)
        except ZeroDivisionError:
            delta_traffic = 100

        if method in ("Credit Card", "Debit Card", "CC emi", "DC emi"):
            print(f"        {Counter2}. For {inner_method} - {inner_aggregator}, {name} - {merchant_id}, network_tokenised_payment - {network_tokenised_payment}, {network}, sr dropped by {vmm_tx} due to change in volume by {delta_traffic}. Previous attempts {previous_attempts} and Current attempts - {current_attempts }")

        if method in ("UPI Intent", "UPI Collect"):
            print(f"        {Counter2}. For {inner_method} - {inner_aggregator}, {name} - {merchant_id}, sr dropped by {vmm_tx} due to change in volume by {delta_traffic}. Previous attempts {previous_attempts} and Current attempts - {current_attempts }")

        if method in ("netbanking"):
            print(f"        {Counter2}. For {inner_method} - {inner_aggregator}, {name} - {merchant_id}, {bank}, sr dropped by {vmm_tx} due to change in volume by {delta_traffic}. Previous attempts {previous_attempts} and Current attempts - {current_attempts }")

        if method in ("wallet"):
            print(f"        {Counter2}. For {inner_method} - {inner_aggregator}, {name} - {merchant_id}, sr dropped by  {vmm_tx} due to change in volume by {delta_traffic}.Previous attempts {previous_attempts} and Current attempts - {current_attempts }")
        Counter2+=1
    print()

# COMMAND ----------

print(method_SR.columns)

# COMMAND ----------


def add_summary_bullet_points(doc, summary_text):
    # Add a heading for the summary section

    # Add bullet points for the summary text
    for text in summary_text:
        # Set font to Muli
        paragraph = doc.add_paragraph(text, style='ListBullet')
        for run in paragraph.runs:
            run.font.name = "Muli"
            run.font.size = Pt(11)  # Set font size to 10

def add_dataframe_as_table(doc, df, name):
    # Add a heading for the DataFrame
    doc.add_heading(name, level=2)

    # Add table
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

    # Set column widths
    table.width = Inches(15)
    column_widths = [1.75] * df.shape[1]  # Set each column width to 1.5 inches
    for column, width in zip(table.columns, column_widths):
        column.width = Inches(width)

    # Set table style
    table.style = "Table Grid"

    # Set borders to black
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)  # Black

    for j in range(df.shape[1]):
        table.cell(0, j).text = df.columns[j]
        # Access paragraph in the cell
        paragraph = table.cell(0, j).paragraphs[0]
        paragraph.runs[0].font.bold = True  # Set first row bold
        paragraph.runs[0].font.name = "Muli"  # Set font to Muli
        paragraph.runs[0].font.size = Pt(10)  # Set font size to 10

    # Add data
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            table.cell(i + 1, j).text = str(df.values[i, j])
            # Access paragraph in the cell
            paragraph = table.cell(i + 1, j).paragraphs[0]
            paragraph.runs[0].font.name = "Muli"  # Set font to Muli
            paragraph.runs[0].font.size = Pt(10)  # Set font size to 10

# Create a new Word document
doc = Document()

for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.name = "Muli"

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Muli"

# Add heading with font Lato and size 30
heading_text = "Optimizer - SR insights for Apollo"
current_date = datetime.now().strftime("%Y-%m-%d")
heading_text += current_date
heading = doc.add_heading(heading_text, level=1)
heading.runs[0].font.name = "Muli"
heading.runs[0].font.size = Pt(30)

doc.add_paragraph(f"Note:")
intro_bullet = [f"T1 period starts on {start_date} & T1 period ends on {end_date}",
f"T2 period starts on {previous_start_date} & T2 period ends on {previous_end_date}"]
add_summary_bullet_points(doc, intro_bullet)

heading_SR_summary = f"Overall Success Rate Delta:{delta_SR_overall}"
heading = doc.add_heading(heading_SR_summary, level=3)
heading.runs[0].font.name = "Muli"
heading.runs[0].font.size = Pt(17)

heading_SR = f"Current month SR:{SR_T1_overall}, Previous month SR:{SR_T2_overall}"
heading = doc.add_heading(heading_SR, level=4)
heading.runs[0].font.name = "Muli"
heading.runs[0].font.size = Pt(14)


selected_columns = method_SR.loc[:,[ 'method_advanced','SR_T-1','SR_T-2','delta_SR']]
        
add_dataframe_as_table(doc, selected_columns,"Method level SR delta")



# Save Word document to local file system with timestamp appended to the file name
current_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
local_output_file_path = f"/tmp/output_{current_time}.docx"
# local_output_file_path = f"/tmp/Post_Pod_AWT_{current_time}.docx"  # Updated file name
doc.save(local_output_file_path)

# Move file to DBFS location
dbfs_output_file_path = f"/dbfs/FileStore/Chitraxi/output.docx"
shutil.move(local_output_file_path, dbfs_output_file_path)
display(HTML(f'<a href="{dbfs_output_file_path}" download>SR insights</a>'))

# COMMAND ----------

#spark.read.format('docx').load('/dbfs/FileStore/Chitraxi/output.docx').display()

# COMMAND ----------



top5_vmm_tx.to_csv("/dbfs/FileStore/Chitraxi/SR_filtered.csv")


spark.read.format('csv').load('dbfs:/FileStore/Chitraxi/top5_vmm_tx.csv').display()
